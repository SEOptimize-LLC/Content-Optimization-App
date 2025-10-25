"""
Semantic SEO Content Optimizer - V2
4-Step Workflow implementing Koray's Semantic SEO Framework
"""

import streamlit as st
import io
import traceback
from docx import Document

# Import semantic SEO optimizers
try:
    # Use NEW semantic extractor V2 (proper semantic analysis)
    from core.semantic_extractor_v2 import extract_entity_context
    from core.outline_optimizer import optimize_outline
    from core.draft_optimizer import optimize_draft
    from core.llm_config import get_llm_client, is_llm_enabled
    print("SUCCESS: Semantic SEO optimizers loaded (using V2 extractor)")
except Exception as e:
    print(f"ERROR loading optimizers: {e}")
    traceback.print_exc()

# Page config
st.set_page_config(
    page_title="Semantic SEO Optimizer",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize session state
if 'step' not in st.session_state:
    st.session_state.step = 1
if 'entity_extraction' not in st.session_state:
    st.session_state.entity_extraction = None
if 'central_entity' not in st.session_state:
    st.session_state.central_entity = None
if 'source_context' not in st.session_state:
    st.session_state.source_context = None
if 'query_report_text' not in st.session_state:
    st.session_state.query_report_text = None
if 'optimization_result' not in st.session_state:
    st.session_state.optimization_result = None
if 'output_md' not in st.session_state:
    st.session_state.output_md = ""
if 'metadata' not in st.session_state:
    st.session_state.metadata = {}

# Title and header
st.title("🎯 Semantic SEO Content Optimizer")
st.markdown("*Based on Koray Tuğberk GÜBÜR's Semantic SEO Framework*")
st.markdown("---")

# Sidebar - Progress tracker
with st.sidebar:
    st.header("📊 Progress")

    steps = [
        "1️⃣ Upload Query Report",
        "2️⃣ Confirm Entity & Context",
        "3️⃣ Select Optimization Mode",
        "4️⃣ Generate & Export"
    ]

    for i, step_name in enumerate(steps, 1):
        if i < st.session_state.step:
            st.success(f"✓ {step_name}")
        elif i == st.session_state.step:
            st.info(f"→ {step_name}")
        else:
            st.text(f"  {step_name}")

    st.markdown("---")

    # Reset button
    if st.button("🔄 Start Over"):
        for key in ['step', 'entity_extraction', 'central_entity', 'source_context',
                    'query_report_text', 'optimization_result', 'output_md', 'metadata']:
            st.session_state[key] = None if key != 'step' else 1
            if key in ['output_md']:
                st.session_state[key] = ""
            if key in ['metadata']:
                st.session_state[key] = {}
        st.rerun()

    st.markdown("---")
    st.markdown("### 🤖 LLM Enhancement")

    # Check LLM status
    llm_available = is_llm_enabled()

    if llm_available:
        st.success("✓ LLM API Configured")
        llm_client = get_llm_client()
        if llm_client:
            st.markdown(f"**Total Cost:** ${llm_client.total_cost:.4f}")
    else:
        st.warning("⚠️ LLM Not Configured")
        st.markdown("Add `OPENROUTER_API_KEY` to Streamlit Secrets to enable AI-enhanced content generation.")

    st.markdown("---")
    st.markdown("### 📚 Framework Features")
    st.markdown("""
    ✅ Entity & Attribute Detection
    ✅ Topical Map Generation
    ✅ Query Processing
    ✅ Content Brief Creation
    ✅ Distributional Semantics
    ✅ Semantic Content Network
    ✅ URL Structure Planning
    ✅ Meta Optimization
    ✅ LLM-Enhanced Definitions
    ✅ AI Keyword Insertion
    """)

# STEP 1: Upload Query Fan-Out Report
if st.session_state.step == 1:
    st.header("Step 1: Upload Query Fan-Out Report")
    st.markdown("Upload your Query Fan-Out report and we'll automatically detect entities and attributes.")

    col1, col2 = st.columns([2, 1])

    with col1:
        query_report = st.file_uploader(
            "Upload Query Fan-Out Report",
            type=['md', 'txt'],
            key="query_report_upload"
        )

        if query_report:
            query_text = query_report.read().decode("utf-8")
            st.session_state.query_report_text = query_text

            # Show preview
            with st.expander("📄 Preview Report"):
                st.text(query_text[:500] + "..." if len(query_text) > 500 else query_text)

            if st.button("🔍 Analyze Report", type="primary"):
                with st.spinner("Analyzing queries and detecting entities..."):
                    try:
                        # Extract entities and context
                        entity_extraction = extract_entity_context(query_text)
                        st.session_state.entity_extraction = entity_extraction
                        st.session_state.step = 2
                        st.success("✓ Analysis complete! Proceed to Step 2.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error during analysis: {str(e)}")
                        st.code(traceback.format_exc())

    with col2:
        st.info("""
        **What is a Query Fan-Out Report?**

        A Query Fan-Out report contains:
        - Related search queries
        - Sub-queries
        - Search insights
        - User intent patterns

        Formats: Markdown (.md) or Text (.txt)
        """)

# STEP 2: Confirm Entity & Source Context
elif st.session_state.step == 2:
    st.header("Step 2: Confirm Central Entity & Source Context")
    st.markdown("Review our auto-detected suggestions and confirm or customize.")

    if st.session_state.entity_extraction:
        extraction = st.session_state.entity_extraction

        col1, col2 = st.columns(2)

        with col1:
            st.subheader("🎯 Central Entity")
            st.markdown("*The main topic/entity your content focuses on*")

            # Show top 5 entity suggestions
            entity_suggestions = extraction.get('entity_suggestions', [])

            if entity_suggestions:
                # Create radio options
                entity_options = [f"{s['entity']} (confidence: {s['confidence']:.2%})" for s in entity_suggestions[:5]]
                entity_options.append("Other (specify below)")

                selected_entity_option = st.radio(
                    "Select Central Entity:",
                    entity_options,
                    key="entity_selection"
                )

                # If "Other" selected, show text input
                if "Other" in selected_entity_option:
                    custom_entity = st.text_input("Enter custom entity:", key="custom_entity")
                    final_entity = custom_entity if custom_entity else entity_suggestions[0]['entity']
                else:
                    # Extract entity from selected option
                    final_entity = selected_entity_option.split(" (confidence")[0]

                st.session_state.central_entity = final_entity

                # Show entity details
                with st.expander("📊 Entity Detection Details"):
                    st.markdown(f"**Total Entities Found:** {len(extraction.get('all_entities', []))}")
                    st.markdown(f"**Total Queries Analyzed:** {extraction.get('total_queries', 0)}")
            else:
                st.warning("No entities detected. Please enter manually.")
                manual_entity = st.text_input("Central Entity:", key="manual_entity")
                st.session_state.central_entity = manual_entity

        with col2:
            st.subheader("💼 Source Context")
            st.markdown("*Who you are and what you offer*")

            # Show source context suggestions
            context_suggestions = extraction.get('source_context_suggestions', [])

            if context_suggestions:
                selected_context = st.selectbox(
                    "Select Source Context:",
                    context_suggestions + ["Other (specify below)"],
                    key="context_selection"
                )

                if selected_context == "Other (specify below)":
                    custom_context = st.text_area(
                        "Describe your source context:",
                        placeholder="E.g., 'Luxury eyewear brand selling designer glasses'",
                        key="custom_context"
                    )
                    final_context = custom_context if custom_context else context_suggestions[0]
                else:
                    final_context = selected_context

                st.session_state.source_context = final_context
            else:
                st.warning("Could not infer source context. Please enter manually.")
                manual_context = st.text_area(
                    "Source Context:",
                    placeholder="E.g., 'Visa consultancy helping people move to Germany'",
                    key="manual_context"
                )
                st.session_state.source_context = manual_context

        # Navigation
        st.markdown("---")
        col_nav1, col_nav2 = st.columns([1, 1])

        with col_nav1:
            if st.button("← Back to Step 1"):
                st.session_state.step = 1
                st.rerun()

        with col_nav2:
            if st.button("Proceed to Step 3 →", type="primary"):
                if st.session_state.central_entity and st.session_state.source_context:
                    st.session_state.step = 3
                    st.rerun()
                else:
                    st.error("Please confirm both Central Entity and Source Context")

# STEP 3: Select Optimization Mode
elif st.session_state.step == 3:
    st.header("Step 3: Select Optimization Mode")

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("📝 Outline Optimization")
        st.markdown("""
        **What it does:**
        - Generates complete topical map
        - Creates content briefs with 4 contextual elements
        - Plans semantic content network
        - Optimizes URL structure
        - Generates meta tags

        **Upload:**
        - Original outline (Markdown)
        """)

        outline_file = st.file_uploader("Upload Original Outline", type=['md'], key="outline_upload")

        if outline_file and st.button("🚀 Optimize Outline", type="primary"):
            outline_text = outline_file.read().decode("utf-8")

            with st.spinner("Running complete semantic SEO optimization..."):
                try:
                    enhanced_md, metadata = optimize_outline(
                        st.session_state.query_report_text,
                        outline_text,
                        st.session_state.central_entity,
                        st.session_state.source_context
                    )

                    st.session_state.output_md = enhanced_md
                    st.session_state.metadata = metadata
                    st.session_state.step = 4
                    st.success("✓ Optimization complete!")
                    st.rerun()

                except Exception as e:
                    st.error(f"Error during optimization: {str(e)}")
                    st.code(traceback.format_exc())

    with col2:
        st.subheader("✍️ Draft Optimization")
        st.markdown("""
        **What it does:**
        - Analyzes distributional semantics
        - Keyword co-occurrence optimization
        - Word sequence recommendations
        - Contextual dance planning
        - Anchor segment suggestions

        **Upload:**
        - Original draft (Text/Markdown)
        - Keyword list (CSV/Text)
        """)

        draft_file = st.file_uploader("Upload Original Draft", type=['txt', 'md'], key="draft_upload")
        keyword_file = st.file_uploader("Upload Keyword List", type=['csv', 'txt', 'md'], key="keyword_upload")

        if draft_file and keyword_file and st.button("🚀 Optimize Draft", type="primary"):
            draft_text = draft_file.read().decode("utf-8")
            keyword_text = keyword_file.read().decode("utf-8")

            # Parse keywords
            keywords = []
            for line in keyword_text.split('\n'):
                line = line.strip()
                if line and not line.startswith('#'):
                    # Handle CSV or plain text
                    if ',' in line:
                        keywords.append(line.split(',')[0].strip())
                    else:
                        keywords.append(line.strip('- ').strip())

            keywords = list(set(keywords))  # Deduplicate

            with st.spinner("Running distributional semantics analysis..."):
                try:
                    report_md, metadata = optimize_draft(
                        draft_text,
                        keywords,
                        st.session_state.central_entity,
                        "Your macro context",  # Can be enhanced
                        keywords[:3]  # Top 3 as macro keywords
                    )

                    st.session_state.output_md = report_md
                    st.session_state.metadata = metadata
                    st.session_state.step = 4
                    st.success("✓ Optimization complete!")
                    st.rerun()

                except Exception as e:
                    st.error(f"Error during optimization: {str(e)}")
                    st.code(traceback.format_exc())

    # Navigation
    st.markdown("---")
    if st.button("← Back to Step 2"):
        st.session_state.step = 2
        st.rerun()

# STEP 4: Results & Export
elif st.session_state.step == 4:
    st.header("Step 4: Results & Export")

    # Display metadata
    with st.sidebar:
        st.subheader("📈 Optimization Summary")
        for key, value in st.session_state.metadata.items():
            if isinstance(value, list):
                st.write(f"**{key}:** {len(value)} items")
            else:
                st.write(f"**{key}:** {value}")

    # Display results
    st.markdown("### Optimization Results")

    # Add prominent download buttons at the top
    st.info("📥 **Ready to download?** Use the buttons below or switch to the Export tab.")

    col_top1, col_top2, col_top3 = st.columns([1, 1, 2])

    with col_top1:
        st.download_button(
            label="📄 Download MD",
            data=st.session_state.output_md,
            file_name="semantic_seo_optimization.md",
            mime="text/markdown",
            use_container_width=True
        )

    with col_top2:
        try:
            docx_buffer = io.BytesIO()
            doc = Document()
            doc.add_heading("Semantic SEO Optimization", 0)

            # Add metadata
            for key, value in st.session_state.metadata.items():
                doc.add_paragraph(f"{key}: {value}")
            doc.add_page_break()

            # Add content
            lines = st.session_state.output_md.split('\n')
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith('- '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                elif line.strip():
                    doc.add_paragraph(line)

            doc.save(docx_buffer)
            docx_buffer.seek(0)

            st.download_button(
                label="📝 Download Word",
                data=docx_buffer,
                file_name="semantic_seo_optimization.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        except Exception as e:
            st.error(f"Word export error: {e}")

    st.markdown("---")

    # Display full report (no tabs - export buttons are already at top)
    st.markdown("### 📄 Full Report")
    st.markdown(st.session_state.output_md)

    # Navigation
    st.markdown("---")
    col_nav1, col_nav2 = st.columns([1, 1])

    with col_nav1:
        if st.button("🔄 Start New Optimization"):
            for key in ['step', 'entity_extraction', 'central_entity', 'source_context',
                        'query_report_text', 'optimization_result', 'output_md', 'metadata']:
                st.session_state[key] = None if key != 'step' else 1
                if key in ['output_md']:
                    st.session_state[key] = ""
                if key in ['metadata']:
                    st.session_state[key] = {}
            st.rerun()

    with col_nav2:
        if st.button("← Back to Step 3"):
            st.session_state.step = 3
            st.rerun()

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666;'>
    Built with Streamlit | Powered by Koray's Semantic SEO Framework<br>
    🤖 <a href='https://claude.com/claude-code'>Generated with Claude Code</a>
</div>
""", unsafe_allow_html=True)
