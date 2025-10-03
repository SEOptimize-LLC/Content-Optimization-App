from docx import Document
from docx.shared import Inches
import io
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from google.oauth2.credentials import Credentials
import os

SCOPES = ['https://www.googleapis.com/auth/documents.create']

def export_to_word(content_md, metadata, filename="optimized_output.docx"):
    """
    Export Markdown content and metadata to Word document.
    
    Args:
        content_md (str): Markdown content.
        metadata (dict): Optimization metadata.
        filename (str): Output filename.
    
    Returns:
        BytesIO: Buffer with .docx file.
    """
    docx_buffer = io.BytesIO()
    doc = Document()
    
    # Add title
    doc.add_heading("Blog Optimizer Output", 0)
    
    # Add metadata
    doc.add_heading("Metadata", level=1)
    for key, value in metadata.items():
        doc.add_paragraph(f"{key}: {value}")
    doc.add_page_break()
    
    # Convert Markdown to paragraphs/headings
    lines = content_md.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
    
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

def export_to_google_docs(content_md, metadata, title="Optimized Blog Content"):
    """
    Export to Google Docs using API. Requires credentials.json in project root.
    
    Args:
        content_md (str): Markdown content.
        metadata (dict): Optimization metadata.
        title (str): Document title.
    
    Returns:
        str: Document URL if successful.
    """
    creds = None
    token_path = 'token.json'
    
    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file(token_path, SCOPES)
    
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file('credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        
        with open(token_path, 'w') as token:
            token.write(creds.to_json())
    
    service = build('docs', 'v1', credentials=creds)
    
    # Create document
    document = {
        'title': title,
        'body': {
            'content': [
                {
                    'element': {
                        'textRun': {
                            'content': f"Blog Optimizer Output\n\nMetadata:\n"
                        }
                    }
                }
            ]
        }
    }
    
    # Add metadata
    for key, value in metadata.items():
        document['body']['content'].append({
            'element': {
                'textRun': {
                    'content': f"{key}: {value}\n"
                }
            }
        })
    
    # Simple Markdown to Docs body (basic, no full parsing)
    lines = content_md.split('\n')
    for line in lines:
        line = line.strip()
        if line.startswith('# '):
            document['body']['content'].append({
                'element': {
                    'textRun': {
                        'content': f"\n{line}\n"
                    }
                }
            })
        elif line.startswith('## '):
            document['body']['content'].append({
                'element': {
                    'textRun': {
                        'content': f"\n## {line[3:]}\n"
                    }
                }
            })
        elif line.startswith('- '):
            document['body']['content'].append({
                'element': {
                    'textRun': {
                        'content': f"• {line[2:]}\n"
                    }
                }
            })
        elif line:
            document['body']['content'].append({
                'element': {
                    'textRun': {
                        'content': f"{line}\n"
                    }
                }
            })
    
    request = service.documents().create(body=document)
    response = request.execute()
    
    doc_id = response.get('documentId')
    doc_url = f"https://docs.google.com/document/d/{doc_id}/edit"
    
    return doc_url